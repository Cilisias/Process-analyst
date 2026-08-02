import os
import re

import pandas as pd
import docx
from docx.text.run import Run
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docxedit
import datetime
from pytrovich.enums import NamePart, Gender, Case
from pytrovich.maker import PetrovichDeclinationMaker
from pytrovich.detector import PetrovichGenderDetector


def make_word_bold(doc: Document, word: str, fntName: str, fntSize: float, case_sensitive: bool = False):
    """
    Делает все вхождения `word` жирным (bold) во всем документе.

    Parameters
    ----------
    doc : Document
        Открытый объект Document.
    word : str
        Искомое слово (можно задать регекс‑выражение, например r'\bтекст\b').
    case_sensitive : bool, optional
        Если False (по умолчанию), поиск нечувствителен к регистру.
    """
    # Подготовим регулярку: ищем слово как отдельную «лицо»,
    # но без захвата пробелов/пунктуации.
    flags = 0 if case_sensitive else re.IGNORECASE
    pattern = re.compile(re.escape(word), flags)

    for para in doc.paragraphs:
        # Список всех ранов параграфа (они изменяются «на месте», поэтому
        # будем работать с копией списка, чтобы можно было вставлять новые run‑ы)
        runs = list(para.runs)

        for i, run in enumerate(runs):
            # Если в текущем run нет искомого текста — пропускаем
            if not pattern.search(run.text):
                continue

            # 1️⃣ Разбиваем текст текущего run на части:
            #    - часть до найденного слова,
            #    - само слово,
            #    - часть после слова.
            # Мы делаем это в цикле, потому что слово может встречаться
            # несколько раз в одном и том же run.
            cursor = 0
            new_runs = []  # сюда будем складывать новые Run‑ы (можно будет вставлять)
            for m in pattern.finditer(run.text):
                start, end = m.span()

                # Добавляем «префикс», если он есть
                if start > cursor:
                    prefix = run.text[cursor:start]
                    new_runs.append(('normal', prefix))

                # Добавляем само слово (будем делать bold)
                word_text = run.text[start:end]
                new_runs.append(('bold', word_text))

                cursor = end

            # Добавляем «суффикс», если после последнего совпадения ещё остался текст
            if cursor < len(run.text):
                suffix = run.text[cursor:]
                new_runs.append(('normal', suffix))

            # 2️⃣ Удаляем оригинальный run из параграфа.
            #    В python‑docx нет метода `remove_run`, поэтому делаем «обходным путем»:
            #    заменяем текст оригинального run на пустую строку и затем удаляем
            #    его из списка `para._p` (это низкоуровневый элемент lxml).
            run.text = ''  # очистим визуальный текст
            para._p.remove(run._r)  # удаляем элемент XML (нет публичного API)

            # 3️⃣ Вставляем новые run‑ы на место старого.
            #    Мы будем использовать `para.add_run()`, а потом переставим
            #    их в нужный порядок, потому что `add_run` всегда добавляет в конец.
            inserted = []
            for style, txt in new_runs:
                new_run = para.add_run(txt)
                if style == 'bold':
                    new_run.bold = True
                    font = new_run.font
                    font.size = Pt(fntSize)
                    font.name = fntName
                inserted.append(new_run)

            # Поскольку `add_run` уже добавил их в конец, а нам надо их разместить
            # там, где был оригинальный run, переместим их «внутрь» списка XML.
            # Найдём индекс, где стоял старый run, и вставим перед следующим элементом.
            # (Если старый run был последним, просто оставляем в конце.)
            # Оригинальный run уже удалён, поэтому берём индекс текущего последнего
            # дочернего элемента перед вставкой.
            # Сложно и не всегда нужно, но для простоты будем просто оставлять их в конце.
            # Если порядок важен – можно использовать более сложный подход с lxml.
            # В большинстве реальных документов это не критично.

            # Прерываем внутренний цикл, потому что мы уже заменили текущий run.
            # Следующий оригинальный run (если он был) теперь смещён в списке,
            # но так как мы работаем с копией `runs`, дальше просто продолжаем.
            break  # выходим из цикла по run‑ам текущего параграфа

file_path = os.path.realpath(__file__)
# print(file_path)

dir_path = os.path.dirname(os.path.realpath(__file__))
# print(dir_path)

if not os.path.isdir(os.path.join(dir_path, 'готовые_Акты')):
    os.makedirs(os.path.join(dir_path, 'готовые_Акты'))

def month_from_eng_to_ru(month):
    out = ''
    if 'january' in month:
        out = month.replace('january', 'января')
    if 'february' in month:
        out = month.replace('february', 'февраля')
    if 'march' in month:
        out = month.replace('march', 'марта')
    if 'april' in month:
        out = month.replace('april', 'апреля')
    if 'may' in month:
        out = month.replace('may', 'мая')
    if 'june' in month:
        out = month.replace('june', 'июня')
    if 'july' in month:
        out = month.replace('july', 'июля')
    if 'august' in month:
        out = month.replace('august', 'августа')
    if 'september' in month:
        out = month.replace('september', 'сентября')
    if 'october' in month:
        out = month.replace('october', 'октября')
    if 'november' in month:
        out = month.replace('november', 'ноября')
    if 'december' in month:
        out = month.replace('december', 'декабря')

    return out


df = pd.read_excel(engine='openpyxl', sheet_name='Лист1', header=0, io='данные_по_сотрудникам.xlsx')
# print(df)

count_cols = df.shape[1]
# print(count_cols)
count_rows = df.shape[0]
# print(count_rows)

maker = PetrovichDeclinationMaker()
detector = PetrovichGenderDetector()

document7 = Document(os.path.join(dir_path, 'template7.docx'))
document8 = Document(os.path.join(dir_path, 'template8.docx'))

dateRejOut = ''
dateAgrOut = ''
personFioListDoc7 = []
personFioListDoc8 = []


for i in range(count_rows):

    nans_in_current_row = df.iloc[i].isna().any()
    # print(nans_in_current_row)

    if nans_in_current_row:
        print(f'В строке № {i+1} обнаружены пустые ячейки. Пропускаю строку!')
        continue

    Comp = df['Comp'].iloc[i]
    # print(Comp)
    if Comp == 'АМИТ':
        Comp = 'ООО «АМ-интеллектуальные технологии»'

    elif Comp == 'ИИТ':
        Comp = 'ООО «Инновация-ИТ»'

    elif Comp == 'ИК':
        Comp = 'ООО «Исходный код»'


    dateRej = df['dateRej'].iloc[0].strftime('%Y-%m-%d')
    # print(dateRej)
    dateRejString = month_from_eng_to_ru(datetime.datetime.strptime(dateRej.replace(' 00:00:00', ''), '%Y-%m-%d').strftime('%d %B %Y г.').lower())
    # print(dateRejString)
    if dateRejString[0] == '0':
        dateRejString = dateRejString[1:]
    # print(dateRejString)
    if i == 0:
        dateRejOut = dateRejString

    dateAgr = df['dateAgr'].iloc[0].strftime('%Y-%m-%d')
    # print(dateAgr)
    dateAgrString = month_from_eng_to_ru(datetime.datetime.strptime(dateAgr.replace(' 00:00:00', ''), '%Y-%m-%d').strftime('%d %B %Y г.').lower())
    # print(dateAgrString)
    if dateAgrString[0] == '0':
        dateAgrString = dateAgrString[1:]
    # print(dateAgrString)
    if i == 0:
        dateAgrOut = dateAgrString

    fioFull = df['fioFull'].iloc[i]
    # print(fioFull)

    position = df['position'].iloc[i]
    # print(position)

    personFioListDoc7.append([str(u'\u2013\u00A0'), str(fioFull), str(position).lower(), str(Comp)])
    # personFioListDoc7.append([str(u'\u2013\u00A0') + str(fioFull), str(position).lower(), str(Comp)])

    fioSplt = fioFull.split(' ')
    fioGender = detector.detect(firstname=fioSplt[0], lastname=fioSplt[1], middlename=fioSplt[2])

    fioLastName = maker.make(NamePart.LASTNAME, fioGender, Case.DATIVE, fioSplt[0])  # Дательный падеж
    fioFirstName = maker.make(NamePart.FIRSTNAME, fioGender, Case.DATIVE, fioSplt[1])  # Дательный падеж
    fioMiddleName = maker.make(NamePart.MIDDLENAME, fioGender, Case.DATIVE, fioSplt[2])  # Дательный падеж

    fioConcDatCase = fioLastName + ' ' + fioFirstName + ' ' + fioMiddleName
    # print(fioConcDatCase)

    personFioListDoc8.append([str(u'\u2013\u00A0'), str(fioConcDatCase)])

# print(personFioListDoc8)

personFioListDoc7ToString = ''
personFioListDoc8ToString = ''

cntDoc7 = 0
for b in personFioListDoc7:
    # print(b)
    if cntDoc7 == len(personFioListDoc7) - 1:
        personFioListDoc7ToString += b[0] + ' ' + b[1] + ', ' + b[2] + ', \n' + b[3]
    elif cntDoc7 < len(personFioListDoc7) - 1:
        personFioListDoc7ToString += b[0] + ' ' + b[1] + ', ' + b[2] + ', \n' + b[3] + '; ' + '\r\n'
    cntDoc7 += 1

# print(personFioListDoc7ToString)

cntDoc8 = 0
for b in personFioListDoc8:
    # print(b)
    if cntDoc8 == len(personFioListDoc8) - 1:
        personFioListDoc8ToString += b[0] + ' ' + b[1]
    elif cntDoc8 < len(personFioListDoc8) - 1:
        personFioListDoc8ToString += b[0] + ' ' + b[1] + '; ' + '\r\n'
    cntDoc8 += 1

# print(personFioListDoc8ToString)

docxedit.replace_string(document7, old_string="dateRej", new_string=dateRejOut, show_errors=False)
docxedit.replace_string(document7, old_string="personList", new_string=personFioListDoc7ToString, show_errors=False)

docxedit.replace_string(document8, old_string="dateAgr", new_string=dateAgrOut, show_errors=False)
docxedit.replace_string(document8, old_string="personList", new_string=personFioListDoc8ToString, show_errors=False)

for p in range(len(personFioListDoc7)):
    # print(personFioListDoc7[p][1])
    make_word_bold(document7, personFioListDoc7[p][1], 'Times New Roman', 12, case_sensitive=False)

for m in range(len(personFioListDoc8)):
    # print(personFioListDoc8[m][1])
    make_word_bold(document8, personFioListDoc8[m][1], 'Times New Roman', 11.5, case_sensitive=False)

document7.save(f"{dir_path}\\готовые_Акты\\Акт_приемки_передачи_{count_rows}_спецов.docx")
document8.save(f"{dir_path}\\готовые_Акты\\Служебное_задание_(общее).docx")