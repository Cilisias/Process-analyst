import os
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docxedit
import datetime
from pytrovich.enums import NamePart, Gender, Case
from pytrovich.maker import PetrovichDeclinationMaker
from pytrovich.detector import PetrovichGenderDetector

file_path = os.path.realpath(__file__)
# print(file_path)

dir_path = os.path.dirname(os.path.realpath(__file__))
# print(dir_path)

if not os.path.isdir(os.path.join(dir_path, 'готовые_Акты')):
    os.makedirs(os.path.join(dir_path, 'готовые_Акты'))

def month_from_eng_to_ru(month):
    out = ''
    if 'january' in month:
        out = month.replace('january', 'января')
    if 'february' in month:
        out = month.replace('february', 'февраля')
    if 'march' in month:
        out = month.replace('march', 'марта')
    if 'april' in month:
        out = month.replace('april', 'апреля')
    if 'may' in month:
        out = month.replace('may', 'мая')
    if 'june' in month:
        out = month.replace('june', 'июня')
    if 'july' in month:
        out = month.replace('july', 'июля')
    if 'august' in month:
        out = month.replace('august', 'августа')
    if 'september' in month:
        out = month.replace('september', 'сентября')
    if 'october' in month:
        out = month.replace('october', 'октября')
    if 'november' in month:
        out = month.replace('november', 'ноября')
    if 'december' in month:
        out = month.replace('december', 'декабря')

    return out


df = pd.read_excel(engine='openpyxl', sheet_name='Лист1', header=0, io='данные_по_сотрудникам.xlsx')
# print(df)

count_cols = df.shape[1]
# print(count_cols)
count_rows = df.shape[0]
print(count_rows)

maker = PetrovichDeclinationMaker()
detector = PetrovichGenderDetector()

for i in range(count_rows):

    nans_in_current_row = df.iloc[i].isna().any()
    # print(nans_in_current_row)

    if nans_in_current_row:
        print(f'В строке № {i+1} обнаружены пустые ячейки. Пропускаю строку!')
        continue

    document1 = Document(os.path.join(dir_path, 'template1.docx'))
    # document2 = Document(os.path.join(dir_path, 'template2.docx'))
    # document3 = Document(os.path.join(dir_path, 'template3.docx'))
    document4 = Document(os.path.join(dir_path, 'template4.docx'))
    document5 = Document(os.path.join(dir_path, 'template5.docx'))
    document6 = Document(os.path.join(dir_path, 'template6.docx'))

    numPos = df['numPos'].iloc[i]
    # print(numPos)
    docxedit.replace_string(document5, old_string="numPos", new_string=numPos, show_errors=False)

    Comp = df['Comp'].iloc[i]
    # print(Comp)

    dateAgr = df['dateAgr'].iloc[i].strftime('%Y-%m-%d')
    # print(dateAgr)
    dateAgrString = month_from_eng_to_ru(datetime.datetime.strptime(dateAgr.replace(' 00:00:00', ''), '%Y-%m-%d').strftime('%d %B %Y г.').lower())
    if dateAgrString[0] == '0':
        dateAgrString = dateAgrString[1:]

    # docxedit.replace_string(document2, old_string="dateAgr", new_string=dateAgrString, show_errors=False)
    # docxedit.replace_string(document3, old_string="dateAgr", new_string=dateAgrString, show_errors=False)
    docxedit.replace_string(document4, old_string="dateAgr", new_string=dateAgrString, show_errors=False)
    docxedit.replace_string(document5, old_string="dateAgr", new_string=dateAgrString, show_errors=False)

    dateRej = df['dateRej'].iloc[i].strftime('%Y-%m-%d')
    # print(dateRej)
    dateRejString = month_from_eng_to_ru(datetime.datetime.strptime(dateRej.replace(' 00:00:00', ''), '%Y-%m-%d').strftime('%d %B %Y г.').lower())
    # print(dateRejString)
    if dateRejString[0] == '0':
        dateRejString = dateRejString[1:]
    # print(dateRejString)
    docxedit.replace_string(document1, old_string="dateRej", new_string=dateRejString, show_errors=False)
    docxedit.replace_string(document6, old_string="dateRej", new_string=dateRejString, show_errors=False)
    #
    fioFull = df['fioFull'].iloc[i]
    # print(fioFull)

    fioSplt = fioFull.split(' ')
    fioGender = detector.detect(firstname=fioSplt[0], lastname=fioSplt[1], middlename=fioSplt[2])
    # print(fioGender)

    initialsAndLastName = fioSplt[0] + '_' + fioSplt[1][:1].upper() + '_' + fioSplt[2][:1].upper()
    # print(initialsAndLastName)

    fioLastName = maker.make(NamePart.LASTNAME, fioGender, Case.GENITIVE, fioSplt[0])  # Родительный падеж
    fioFirstName = maker.make(NamePart.FIRSTNAME, fioGender, Case.GENITIVE, fioSplt[1].lower())  # Родительный падеж
    fioMiddleName = maker.make(NamePart.MIDDLENAME, fioGender, Case.GENITIVE, fioSplt[2])  # Родительный падеж

    fioConcGenCase = fioLastName + ' ' + fioFirstName.capitalize() + ' ' + fioMiddleName
    # # print(fioConcGenCase)
    #
    docxedit.replace_string(document1, old_string="fioFull", new_string=fioFull, show_errors=False)
    # docxedit.replace_string(document2, old_string="fioFull", new_string=fioFull, show_errors=False)
    # docxedit.replace_string(document3, old_string="fioFull", new_string=fioFull, show_errors=False)
    docxedit.replace_string(document4, old_string="fioFull", new_string=fioConcGenCase, show_errors=False)
    docxedit.replace_string(document5, old_string="fioFull", new_string=fioFull, show_errors=False)
    docxedit.replace_string(document6, old_string="fioFull", new_string=fioFull, show_errors=False)

    position = df['position'].iloc[i]
    # print(position)
    docxedit.replace_string(document5, old_string="position", new_string=position, show_errors=False)
    docxedit.replace_string(document6, old_string="position", new_string=position, show_errors=False)

    dateBirth = df['dateBirth'].iloc[i].strftime('%d.%m.%Y')
    # print(dateBirth)
    # docxedit.replace_string(document2, old_string="dateBirth", new_string=dateBirth, show_errors=False)
    # docxedit.replace_string(document3, old_string="dateBirth", new_string=dateBirth, show_errors=False)
    docxedit.replace_string(document4, old_string="dateBirth", new_string=dateBirth, show_errors=False)

    numPass = df['numPass'].iloc[i]
    # print(numPass)
    # docxedit.replace_string(document2, old_string="numPass", new_string=numPass, show_errors=False)
    # docxedit.replace_string(document3, old_string="numPass", new_string=numPass, show_errors=False)
    docxedit.replace_string(document4, old_string="numPass", new_string=numPass, show_errors=False)

    issuedPass = df['issuedPass'].iloc[i]
    # print(issuedPass)
    # docxedit.replace_string(document2, old_string="issuedPass", new_string=issuedPass, show_errors=False)
    # docxedit.replace_string(document3, old_string="issuedPass", new_string=issuedPass, show_errors=False)
    docxedit.replace_string(document4, old_string="issuedPass", new_string=issuedPass, show_errors=False)

    dateIss = df['dateIss'].iloc[i].strftime('%d.%m.%Y')
    # print(dateIss)
    # docxedit.replace_string(document2, old_string="dateIss", new_string=dateIss, show_errors=False)
    # docxedit.replace_string(document3, old_string="dateIss", new_string=dateIss, show_errors=False)
    docxedit.replace_string(document4, old_string="dateIss", new_string=dateIss, show_errors=False)

    depCode = df['depCode'].iloc[i]
    # print(depCode)
    # docxedit.replace_string(document2, old_string="depCode", new_string=depCode, show_errors=False)
    # docxedit.replace_string(document3, old_string="depCode", new_string=depCode, show_errors=False)
    docxedit.replace_string(document4, old_string="depCode", new_string=depCode, show_errors=False)

    regAddress = df['regAddress'].iloc[i]

    if fioGender == Gender.FEMALE:
        # docxedit.replace_string(document2, old_string="regAddress", new_string='зарегистрированная по адресу: ' + regAddress, show_errors=False)
        # docxedit.replace_string(document3, old_string="regAddress", new_string='зарегистрированная по адресу: ' + regAddress, show_errors=False)
        docxedit.replace_string(document4, old_string="regAddress", new_string='зарегистрирована по адресу: ' + regAddress, show_errors=False)
    elif fioGender == Gender.MALE:
        # docxedit.replace_string(document2, old_string="regAddress", new_string='зарегистрированный по адресу: ' + regAddress, show_errors=False)
        # docxedit.replace_string(document3, old_string="regAddress", new_string='зарегистрированный по адресу: ' + regAddress, show_errors=False)
        docxedit.replace_string(document4, old_string="regAddress", new_string='зарегистрирован по адресу: ' + regAddress, show_errors=False)

    # print(regAddress)

    document1.save(f"{dir_path}\\готовые_Акты\\Отказ_от_упоминания_{initialsAndLastName}.docx")
    # document2.save(f"{dir_path}\\готовые_Акты\\Согласие_на_обработку_№1_{initialsAndLastName}.docx")
    # document3.save(f"{dir_path}\\готовые_Акты\\Согласие_на_обработку_№2_{initialsAndLastName}.docx")
    document4.save(f"{dir_path}\\готовые_Акты\\Согласие_на_обработку_{initialsAndLastName}.docx")
    document5.save(f"{dir_path}\\готовые_Акты\\Лист_ознакомления_со_Служеб._заданием_{initialsAndLastName}.docx")
    document6.save(f"{dir_path}\\готовые_Акты\\Акт приемки передачи_{initialsAndLastName}.docx")

    # exit()