import dearpygui.dearpygui as dpg
import dearpygui_animate as animate
import pandas as pd
from pytrovich.enums import NamePart, Gender, Case
from pytrovich.maker import PetrovichDeclinationMaker
from pytrovich.detector import PetrovichGenderDetector
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docxedit
import datetime
# import sqlean
from functools import reduce
import operator
import re
from num2words import num2words
import os
import warnings
import sys
from pathlib import Path
import duckdb
from xlrd.xldate import xldate_as_datetime

warnings.filterwarnings("ignore")

# sqlean.extensions.enable_all()

def get_executable_dir():
    if getattr(sys, 'frozen', False):  # Проверяем, запущен ли скрипт как exe (скомпилирован)
        executable_path = sys.executable  # Путь к exe-файлу
        return Path(executable_path).parent.resolve()
    else:
        return Path(__file__).parent.resolve()  # Если скрипт запущен как .py

# scriptDir = os.path.dirname(os.path.realpath(sys.argv[0]))
scriptDir = get_executable_dir()

print(scriptDir)

def month_from_eng_to_ru(month):
    out = ''
    if 'january' in month:
        out = month.replace('january', 'января')
    if 'february' in month:
        out = month.replace('february', 'февраля')
    if 'march' in month:
        out = month.replace('march', 'марта')
    if 'april' in month:
        out = month.replace('april', 'апреля')
    if 'may' in month:
        out = month.replace('may', 'мая')
    if 'june' in month:
        out = month.replace('june', 'июня')
    if 'july' in month:
        out = month.replace('july', 'июля')
    if 'august' in month:
        out = month.replace('august', 'августа')
    if 'september' in month:
        out = month.replace('september', 'сентября')
    if 'october' in month:
        out = month.replace('october', 'октября')
    if 'november' in month:
        out = month.replace('november', 'ноября')
    if 'december' in month:
        out = month.replace('december', 'декабря')

    return out

dpg.create_context()



big_let_start = 0x00C0  # Capital "A" in cyrillic alphabet
big_let_end = 0x00DF  # Capital "Я" in cyrillic alphabet
small_let_end = 0x00FF  # small "я" in cyrillic alphabet
remap_big_let = 0x0410  # Starting number for remapped cyrillic alphabet
alph_len = big_let_end - big_let_start + 1  # adds the shift from big letters to small
alph_shift = remap_big_let - big_let_start  # adds the shift from remapped to non-remapped
with dpg.font_registry():
    with dpg.font("OpenSans-Regular.ttf", 18) as default_font:
        dpg.add_font_range_hint(dpg.mvFontRangeHint_Default)
        dpg.add_font_range_hint(dpg.mvFontRangeHint_Cyrillic)
        biglet = remap_big_let  # Starting number for remapped cyrillic alphabet
        for i1 in range(big_let_start, big_let_end + 1):  # Cycle through big letters in cyrillic alphabet
            dpg.add_char_remap(i1, biglet)  # Remap the big cyrillic letter
            dpg.add_char_remap(i1 + alph_len, biglet + alph_len)  # Remap the small cyrillic letter
            biglet += 1  # choose next letter
        dpg.bind_font(default_font)

def crbutton_callback(sender, app_data, user_data):
    if dpg.get_value('year') == '':
        animate.add("position", "year", [220, 100], [250, 100], [0, .06, .2, .99], 50)
    if dpg.get_value('fio') == '':
        animate.add("position", "fio", [30, 40], [5, 40], [0, .06, .2, .99], 50)

    if dpg.get_value('year') != '' and dpg.get_value('fio') != '':
        # print(f"fio: {dpg.get_value('fio')}")
        # print(f"year: {dpg.get_value('year')}")

        dpg.set_value(item='error', value='')
        dpg.set_viewport_height(500)
        dpg.set_viewport_width(700)
        dpg.show_item(item="file_dialog_id")


def filedialog_ok_callback(sender, app_data, user_data):
    lenOfSelections = len(app_data['selections'])
    if lenOfSelections == 0:
        print("Не выбран файл таблицы")
        dpg.show_item(item="file_dialog_id")

    elif lenOfSelections == 1:
        # print("OK")
        dpg.set_viewport_height(380)
        dpg.set_viewport_width(420)

        fioFromListSplit = dpg.get_value('fio').split('\n')
        # print(fioFromListSplit)

        checkboxStatus = dpg.get_value('new')
        # print(dpg.get_value('new'))

        company = dpg.get_value('comp')

        if not os.path.isdir(f"{scriptDir}\\готовые Акты"):
            os.mkdir(f"{scriptDir}\\готовые Акты")

        positionAtWorkTXT = open('positionAtWork.txt', 'r', encoding='utf8')
        positionAtWorkTXTRead = positionAtWorkTXT.read()
        positionAtWorkTXTSplit = list(filter(None, positionAtWorkTXTRead.strip().replace('\n', '').split(';')))
        positionAtWorkTXT.close()
        # print(positionAtWorkTXTSplit)

        for itemFIO in fioFromListSplit:
            print(itemFIO)

            try:
                fioInDB = duckdb.sql(f"SELECT \"ФИО сотрудника\" "
                                     f"FROM read_xlsx('{app_data['file_path_name']}', header=true, sheet='Свод', all_varchar=true, stop_at_empty=false, ignore_errors=true) "
                                     f"WHERE \"ФИО сотрудника\" like '{itemFIO.strip()}%' "
                                     f"AND \"Год\" = {int(dpg.get_value('year'))} "
                                     f"AND \"Компания\" like '%{str(dpg.get_value('comp'))}%' "
                                     f"{"AND LOWER(\"Акт на РИД\") = 'нет'" if checkboxStatus else ''}"
                                     ).fetchone()

                # print(fioInDB)

            except duckdb.BinderException as e:
                # print(e)
                dpg.set_value(item='error', value='')
                dpg.set_value(item='error', value='Ошибка в запросе.\nЧто-то не так в исходной таблице! >>ФИО сотрудника<<')
                print('Ошибка в запросе. Что-то не так в исходной таблице! >>ФИО сотрудника<<', 'Перехожу к следующему ФИО')
                continue

            try:
                positionInDB = duckdb.sql(f"SELECT DISTINCT \"Должность\" "
                                     f"FROM read_xlsx('{app_data['file_path_name']}', header=true, sheet='Свод', all_varchar=true, stop_at_empty = false, ignore_errors = true) "
                                     f"WHERE \"ФИО сотрудника\" like '{itemFIO.strip()}%' "
                                     f"AND \"Год\" = {int(dpg.get_value('year'))} "
                                     f"AND \"Компания\" like '%{str(dpg.get_value('comp'))}%'"
                                     f"{"AND LOWER(\"Акт на РИД\") = 'нет'" if checkboxStatus else ''}"
                                     ).fetchone()

                # print(positionInDB)

            except duckdb.BinderException as e:
                # print(e)
                dpg.set_value(item='error', value='')
                dpg.set_value(item='error', value='Ошибка в запросе.\nЧто-то не так в исходной таблице! >>Должность<<')
                print('Ошибка в запросе. Что-то не так в исходной таблице! >>Должность<<', 'Перехожу к следующему ФИО')
                continue

            try:
                employmentContractNumInDB = duckdb.sql(f"SELECT DISTINCT \"Трудовой договор (номер)\" "
                                          f"FROM read_xlsx('{app_data['file_path_name']}', header=true, sheet='Свод', all_varchar=true, stop_at_empty = false, ignore_errors = true) "
                                          f"WHERE \"ФИО сотрудника\" like '{itemFIO.strip()}%' "
                                          f"AND \"Год\" = {int(dpg.get_value('year'))} "
                                          f"AND \"Компания\" like '%{str(dpg.get_value('comp'))}%'"
                                          f"{"AND LOWER(\"Акт на РИД\") = 'нет'" if checkboxStatus else ''}"
                                          ).fetchone()

                # print(employmentContractNumInDB)

            except duckdb.BinderException as e:
                # print(e)
                dpg.set_value(item='error', value='')
                dpg.set_value(item='error', value='Ошибка в запросе.\nЧто-то не так в исходной таблице! >>Трудовой договор (номер)<<')
                print('Ошибка в запросе. Что-то не так в исходной таблице! >>Трудовой договор (номер)<<', 'Перехожу к следующему ФИО')
                continue

            try:
                employmentContractDateInDB = duckdb.sql(f"SELECT DISTINCT \"Трудовой договор (дата)\" "
                                          f"FROM read_xlsx('{app_data['file_path_name']}', header=true, sheet='Свод', all_varchar=true, stop_at_empty = false, ignore_errors = true) "
                                          f"WHERE \"ФИО сотрудника\" like '{itemFIO.strip()}%' "
                                          f"AND \"Год\" = {int(dpg.get_value('year'))} "
                                          f"AND \"Компания\" like '%{str(dpg.get_value('comp'))}%'"
                                          f"{"AND LOWER(\"Акт на РИД\") = 'нет'" if checkboxStatus else ''}"
                                          ).fetchone()

                # print(employmentContractDateInDB)

            except duckdb.BinderException as e:
                # print(e)
                dpg.set_value(item='error', value='')
                dpg.set_value(item='error', value='Ошибка в запросе.\nЧто-то не так в исходной таблице! >>Трудовой договор (дата)<<')
                print('Ошибка в запросе. Что-то не так в исходной таблице! >>Трудовой договор (дата)<<', 'Перехожу к следующему ФИО')
                continue

            try:
                dateActRIDInDB = duckdb.sql(f"SELECT DISTINCT \"Дата Акта на РИД\" "
                                          f"FROM read_xlsx('{app_data['file_path_name']}', header=true, sheet='Свод', all_varchar=true, stop_at_empty = false, ignore_errors = true) "
                                          f"WHERE \"ФИО сотрудника\" like '{itemFIO.strip()}%' "
                                          f"AND \"Год\" = {int(dpg.get_value('year'))} "
                                          f"AND \"Компания\" like '%{str(dpg.get_value('comp'))}%'"
                                          f"{"AND LOWER(\"Акт на РИД\") = 'нет'" if checkboxStatus else ''}"
                                          ).fetchall()
                dateActRIDInDBFetched = reduce(operator.concat, ([list(i) for i in dateActRIDInDB]))
                # print(dateActRIDInDBFetched)

            except duckdb.BinderException as e:
                # print(e)
                dpg.set_value(item='error', value='')
                dpg.set_value(item='error', value='Ошибка в запросе.\nЧто-то не так в исходной таблице! >>Дата Акта на РИД<<')
                print('Ошибка в запросе. Что-то не так в исходной таблице! >>Дата Акта на РИД<<', 'Перехожу к следующему ФИО')
                continue

            if dateActRIDInDBFetched[0] is None or dateActRIDInDBFetched[0].lower() == 'нет':
                dpg.set_value(item='error', value='')
                dpg.set_value(item='error', value='Ошибка в запросе.\nВ колонке \"Дата Акта на РИД\" не проставлена дата!')
                print('Ошибка в запросе. В колонке \"Дата Акта на РИД\" не проставлена дата!', 'Перехожу к следующему ФИО')
                continue

            if len(dateActRIDInDBFetched) > 1:
                dpg.set_value(item='error', value='')
                dpg.set_value(item='error', value='Ошибка в запросе.\nВ колонке \"Дата Акта на РИД\" найдено\nболее одной уникальной даты!')
                print('Ошибка в запросе. В колонке \"Дата Акта на РИД\" найдено более одной уникальной даты!', 'Перехожу к следующему ФИО')
                continue

            try:
                contractGKInDB = duckdb.sql(f"SELECT DISTINCT \"Договор/ГК\n(дата/номер)\" "
                                          f"FROM read_xlsx('{app_data['file_path_name']}', header=true, sheet='Свод', all_varchar=true, stop_at_empty = false, ignore_errors = true) "
                                          f"WHERE \"ФИО сотрудника\" like '{itemFIO.strip()}%' "
                                          f"AND \"Год\" = {int(dpg.get_value('year'))} "
                                          f"AND \"Компания\" like '%{str(dpg.get_value('comp'))}%'"
                                          f"{"AND LOWER(\"Акт на РИД\") = 'нет'" if checkboxStatus else ''}"
                                          ).fetchall()
                contractGKInDBFetched = reduce(operator.concat, ([list(i) for i in contractGKInDB]))
                # print(contractGKInDBFetched)

            except duckdb.BinderException as e:
                # print(e)
                dpg.set_value(item='error', value='')
                dpg.set_value(item='error', value='Ошибка в запросе.\nЧто-то не так в исходной таблице! >>Договор/ГК (дата/номер)<<')
                print('Ошибка в запросе. Что-то не так в исходной таблице! >>Договор/ГК (дата/номер)<<', 'Перехожу к следующему ФИО')
                continue

            maker = PetrovichDeclinationMaker()
            detector = PetrovichGenderDetector()

            fioSplt = fioInDB[0].split(' ')

            initialsAndLastName = fioSplt[1][:1].upper() + '.' + fioSplt[2][:1].upper() + '. ' + fioSplt[0]
            # print(initialsAndLastName)

            fioGender = detector.detect(firstname=fioSplt[0], lastname=fioSplt[1], middlename=fioSplt[2])

            fioLastName = maker.make(NamePart.LASTNAME, fioGender, Case.INSTRUMENTAL, fioSplt[0]) # Творительный падеж
            fioFirstName = maker.make(NamePart.FIRSTNAME, fioGender, Case.INSTRUMENTAL, fioSplt[1]) # Творительный падеж
            fioMiddleName = maker.make(NamePart.MIDDLENAME, fioGender, Case.INSTRUMENTAL, fioSplt[2]) # Творительный падеж

            fioConcINSCase = fioLastName + ' ' + fioFirstName + ' ' + fioMiddleName
            # print(fioConcINSCase)

            positionInflectionString = ''
            for pos in positionAtWorkTXTSplit:
                if pos.find(positionInDB[0].strip() + '|') != -1:
                    # print(pos)
                    positionInflectionString = pos.split('|')[1]
                    # print(positionInflectionString)

            dateActRIDString = month_from_eng_to_ru(xldate_as_datetime(float(dateActRIDInDBFetched[0]), 0).strftime('«%d» %B %Y').lower())
            # print(dateActRIDString)

            employmentContractDateInDF = month_from_eng_to_ru(xldate_as_datetime(float(employmentContractDateInDB[0]), 0).strftime('«%d» %B %Y').lower())
            # print(employmentContractDateInDF)

            document  = Document("template.docx")

            docxedit.replace_string(document, old_string="posINS", new_string=positionInflectionString.strip(), show_errors=False)
            docxedit.replace_string(document, old_string="fioINS", new_string=fioConcINSCase, show_errors=False)
            docxedit.replace_string(document, old_string="dateAct", new_string=dateActRIDString, show_errors=False)
            docxedit.replace_string(document, old_string="contrEmp", new_string=employmentContractNumInDB[0], show_errors=False)
            docxedit.replace_string(document, old_string="empContDate", new_string=employmentContractDateInDF, show_errors=False)
            docxedit.replace_string(document, old_string="NOMpos", new_string=positionInDB[0], show_errors=False)
            docxedit.replace_string(document, old_string="initLN", new_string=initialsAndLastName, show_errors=False)
            docxedit.replace_string(document, old_string="amount", new_string=str(format(float(len(contractGKInDBFetched) * 100), '.2f')).replace('.', ','), show_errors=False)
            docxedit.replace_string(document, old_string="ght", new_string=str(num2words(len(contractGKInDBFetched) * 100, lang='ru', to='currency')), show_errors=False)

            if company == 'АМИТ':
                docxedit.replace_string(document, old_string="company", new_string=u"АМ" + u'\u2011' + "интеллектуальные технологии", show_errors=False)
                docxedit.replace_string(document, old_string="topManager", new_string="Фридмана Александра Михайловича", show_errors=False)
                docxedit.replace_string(document, old_string="tpManagerInit", new_string="А.М. Фридман", show_errors=False)
            elif company == 'Инновация-ИТ':
                docxedit.replace_string(document, old_string="company", new_string="«Инновация-ИТ»", show_errors=False)
                docxedit.replace_string(document, old_string="topManager", new_string="Чижова Алексея Викторовича", show_errors=False)
                docxedit.replace_string(document, old_string="tpManagerInit", new_string="А.В. Чижов", show_errors=False)


            countRows = 1
            for item in contractGKInDBFetched:
                try:
                    nameISInDB = duckdb.sql(f"SELECT DISTINCT \"Наименование ИС\" "
                                                f"FROM read_xlsx('{app_data['file_path_name']}', header=true, sheet='Свод', all_varchar=true, stop_at_empty = false, ignore_errors = true) "
                                                f"WHERE \"ФИО сотрудника\" like '{itemFIO.strip()}%' "
                                                f"AND \"Год\" = {int(dpg.get_value('year'))} "
                                                f"AND \"Договор/ГК\n(дата/номер)\" like '%{item}%' "
                                                f"AND \"Компания\" like '%{str(dpg.get_value('comp'))}%'"
                                                f"{"AND LOWER(\"Акт на РИД\") = 'нет'" if checkboxStatus else ''}"
                                                ).fetchall()

                    nameISInDBFetched = ','.join(nameISInDB[0])
                    # print(nameISInDBFetched)

                except duckdb.BinderException as e:
                    # print(e)
                    dpg.set_value(item='error', value='')
                    dpg.set_value(item='error', value='Ошибка в запросе.\nЧто-то не так в исходной таблице! >>Наименование ИС<<')
                    print('Ошибка в запросе. Что-то не так в исходной таблице! >>Наименование ИС<<', 'Перехожу к следующему ФИО')
                    continue

                try:
                    baseOfWorkInDB = duckdb.sql(f"SELECT DISTINCT \"Основание для выполнения работ дата, номер\" "
                                                f"FROM read_xlsx('{app_data['file_path_name']}', header=true, sheet='Свод', all_varchar=true, stop_at_empty = false, ignore_errors = true) "
                                                f"WHERE \"ФИО сотрудника\" like '{itemFIO.strip()}%' "
                                                f"AND \"Год\" = {int(dpg.get_value('year'))} "
                                                f"AND \"Договор/ГК\n(дата/номер)\" like '%{item}%' "
                                                f"AND \"Компания\" like '%{str(dpg.get_value('comp'))}%'"
                                                f"{"AND LOWER(\"Акт на РИД\") = 'нет'" if checkboxStatus else ''}"
                                                ).fetchall()

                    baseOfWorkInDBFetched = reduce(operator.concat, ([list(i) for i in baseOfWorkInDB]))
                    # print(baseOfWorkInDBFetched)

                except duckdb.BinderException as e:
                    # print(e)
                    dpg.set_value(item='error', value='')
                    dpg.set_value(item='error', value='Ошибка в запросе.\nЧто-то не так в исходной таблице! >>Основание для выполнения работ дата, номер<<')
                    print('Ошибка в запросе. Что-то не так в исходной таблице! >>Основание для выполнения работ дата, номер<<', 'Перехожу к следующему ФИО')
                    continue

                try:
                    ridResultInDB = duckdb.sql(f"SELECT DISTINCT \"РИД (результат)/ наименование программы\" "
                                                f"FROM read_xlsx('{app_data['file_path_name']}', header=true, sheet='Свод', all_varchar=true, stop_at_empty = false, ignore_errors = true) "
                                                f"WHERE \"ФИО сотрудника\" like '{itemFIO.strip()}%' "
                                                f"AND \"Год\" = {int(dpg.get_value('year'))} "
                                                f"AND \"Договор/ГК\n(дата/номер)\" like '%{item}%' "
                                                f"AND \"Компания\" like '%{str(dpg.get_value('comp'))}%'"
                                                f"{"AND LOWER(\"Акт на РИД\") = 'нет'" if checkboxStatus else ''}"
                                                ).fetchall()

                    ridResultInDBFetched = reduce(operator.concat, ([list(i) for i in ridResultInDB]))
                    # print(ridResultInDBFetched)

                except duckdb.BinderException as e:
                    # print(e)
                    dpg.set_value(item='error', value='')
                    dpg.set_value(item='error', value='Ошибка в запросе.\nЧто-то не так в исходной таблице! >>РИД (результат)/ наименование программы<<')
                    print('Ошибка в запросе. Что-то не так в исходной таблице! >>РИД (результат)/ наименование программы<<', 'Перехожу к следующему ФИО')
                    continue

                try:
                    actCompletedWorksNumDateInDB = duckdb.sql(f"SELECT DISTINCT \"Номер, дата документа о передаче результата работ\" "
                                                f"FROM read_xlsx('{app_data['file_path_name']}', header=true, sheet='Свод', all_varchar=true, stop_at_empty = false, ignore_errors = true) "
                                                f"WHERE \"ФИО сотрудника\" like '{itemFIO.strip()}%' "
                                                f"AND \"Год\" = {int(dpg.get_value('year'))} "
                                                f"AND \"Договор/ГК\n(дата/номер)\" like '%{item}%' "
                                                f"AND \"Компания\" like '%{str(dpg.get_value('comp'))}%'"
                                                f"{"AND LOWER(\"Акт на РИД\") = 'нет'" if checkboxStatus else ''}"
                                                ).fetchall()

                    actCompletedWorksNumDateInDBFetched = reduce(operator.concat, ([list(i) for i in actCompletedWorksNumDateInDB]))
                    # print(actCompletedWorksNumDateInDBFetched)
                except duckdb.BinderException as e:
                    # print(e)
                    dpg.set_value(item='error', value='')
                    dpg.set_value(item='error', value='Ошибка в запросе.\nЧто-то не так в исходной таблице! >>Номер, дата документа о передаче результата работ<<')
                    print('Ошибка в запросе. Что-то не так в исходной таблице! >>Номер, дата документа о передаче результата работ<<', 'Перехожу к следующему ФИО')
                    continue

                try:
                    projectInDB = duckdb.sql(f"SELECT DISTINCT \"Проект\", \"Примечание (наименование проекта, номер и дата договора с контрагентом Работодателя)\" "
                                                f"FROM read_xlsx('{app_data['file_path_name']}', header=true, sheet='Свод', all_varchar=true, stop_at_empty = false, ignore_errors = true) "
                                                f"WHERE \"ФИО сотрудника\" like '{itemFIO.strip()}%' "
                                                f"AND \"Год\" = {int(dpg.get_value('year'))} "
                                                f"AND \"Договор/ГК\n(дата/номер)\" like '%{item}%' "
                                                f"AND \"Компания\" like '%{str(dpg.get_value('comp'))}%'"
                                                f"{"AND LOWER(\"Акт на РИД\") = 'нет'" if checkboxStatus else ''}"
                                                ).fetchall()

                    projectInDBFetched = reduce(operator.concat, ([list(i) for i in projectInDB]))
                    # print(projectInDBFetched)

                except duckdb.BinderException as e:
                    # print(e)
                    dpg.set_value(item='error', value='')
                    dpg.set_value(item='error', value='Ошибка в запросе.\nЧто-то не так в исходной таблице! >>Проект или Примечание<<')
                    print('Ошибка в запросе. Что-то не так в исходной таблице! >>Проект или Примечание<<', 'Перехожу к следующему ФИО')
                    continue


                if projectInDBFetched[1] is not None and projectInDBFetched[1].find('\"') != -1:
                    projectInDBFetched[1] = re.sub(pattern=r'\"(\S)', repl=r'«\1', string=projectInDBFetched[1])
                    projectInDBFetched[1] = re.sub(pattern=r'(\S)\"', repl=r'\1»', string=projectInDBFetched[1])
                    # print(projectInDBFetched[1])

                document.tables[0].add_row()

                cell0 = document.tables[0].cell(countRows, 0)
                cell0.paragraphs[0].add_run(str(countRows))
                cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                docxedit.add_text_in_table(document.tables[0], row_num=countRows, column_num=1, new_string=nameISInDBFetched, show_errors=False)

                if len(baseOfWorkInDBFetched) > 1:
                    baseOfWorkInDBFetchedJoin = ";\n".join(u'\u2013 ' + str(element) for element in baseOfWorkInDBFetched)
                    docxedit.add_text_in_table(document.tables[0], row_num=countRows, column_num=2, new_string=re.sub(' +', ' ', baseOfWorkInDBFetchedJoin.strip()), show_errors=False)
                elif len(baseOfWorkInDBFetched) == 1:
                    docxedit.add_text_in_table(document.tables[0], row_num=countRows, column_num=2, new_string=re.sub(' +', ' ', baseOfWorkInDBFetched[0].strip()), show_errors=False)

                docxedit.add_text_in_table(document.tables[0], row_num=countRows, column_num=3, new_string=re.sub(' +', ' ', ridResultInDBFetched[0].strip()).replace('\n', ' '), show_errors=False)

                if len(actCompletedWorksNumDateInDBFetched) > 1:
                    actCompletedWorksNumDateInDBFetchedJoin = ";\n".join(u'\u2013 ' + str(element) for element in actCompletedWorksNumDateInDBFetched)
                    docxedit.add_text_in_table(document.tables[0], row_num=countRows, column_num=4, new_string=re.sub(' +', ' ', actCompletedWorksNumDateInDBFetchedJoin.strip()), show_errors=False)
                elif len(actCompletedWorksNumDateInDBFetched) == 1:
                    docxedit.add_text_in_table(document.tables[0], row_num=countRows, column_num=4, new_string=re.sub(' +', ' ', actCompletedWorksNumDateInDBFetched[0].strip()), show_errors=False)

                cell5 = document.tables[0].cell(countRows, 5)
                if projectInDBFetched[1] is None:
                    cell5.paragraphs[0].add_run(re.sub(' +', ' ', projectInDBFetched[0].strip())).bold = True
                elif projectInDBFetched[1] is not None:
                    projectInDBFetchedStriped = re.sub(' +', ' ', projectInDBFetched[1].strip()).replace('•', '').replace('\n', '')
                    projectInDBFetchedStripedSplited = projectInDBFetchedStriped.split(';')
                    projectInDBFetchedStripedSplitedRemoveEmpty = list(filter(None, projectInDBFetchedStripedSplited))

                    cell5 = document.tables[0].cell(countRows, 5)
                    if len(projectInDBFetchedStripedSplitedRemoveEmpty) > 1:
                        projectInDBFetchedStripedSplitedRemoveEmptyJoin = ";\n".join(u'\u2013 ' + str(element) for element in projectInDBFetchedStripedSplitedRemoveEmpty)
                        cell5.paragraphs[0].add_run(projectInDBFetched[0].strip()).bold = True
                        cell5.add_paragraph()
                        cell5.paragraphs[1].add_run(projectInDBFetchedStripedSplitedRemoveEmptyJoin).bold = False

                    elif len(projectInDBFetchedStripedSplitedRemoveEmpty) == 1:
                        cell5.paragraphs[0].add_run(projectInDBFetched[0].strip()).bold = True
                        cell5.add_paragraph()
                        cell5.paragraphs[1].add_run(projectInDBFetchedStripedSplitedRemoveEmpty[0]).bold = False


                countRows += 1

            for row in document.tables[0].rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(10)
                            font.name = 'Times New Roman'

            document.save(f"{scriptDir}\\готовые Акты\\{itemFIO.replace(' ', '_')}_{str(company)}_Акт_передачи_ИП.docx")

        print('Формирование завершено')

def cancel_callback_filedialog(sender, app_data, user_data):
    # print('Cancel')
    dpg.set_viewport_height(380)
    dpg.set_viewport_width(420)

# dpg.show_font_manager()


with dpg.window(tag="primeW", width=420, height=380, modal=False, no_close=True, no_resize=False, no_move=False, autosize=False, pos=[0, 0]):
    dpg.add_text("ФИО")
    dpg.add_text('', pos=[5, 280], parent="primeW", color=[255, 0, 0], tag='error')
    dpg.add_input_text(width=230, height=200, no_horizontal_scroll=True, multiline=True, tag="fio", escape_clears_all=True)
    dpg.add_text("Год", pos=[250, 75])
    dpg.add_text("Компания", pos=[250, 150])
    dpg.add_listbox(items=['АМИТ', 'Инновация-ИТ'], width=100, pos=[250, 175], tag="comp")

    with dpg.tooltip(parent="comp", hide_on_activity=True, delay=0.5):
        dpg.add_text("A tooltip")

    dpg.add_input_text(width=100, height=20, pos=[250, 100], scientific=True, tag="year")
    dpg.add_button(label="Создать акт(ы)", tag='crbutton', callback=crbutton_callback)

    dpg.add_checkbox(label="Только новые", tag="new", pos=[240, 270], default_value=True)

dpg.set_value('fio', "")
dpg.set_value('year', "2025")

animate.add("position", "year", [15, 15], [250, 100], [0, .06, .2, .99], 50)
animate.add("position", "comp", [15, 15], [250, 175], [0, .06, .2, .99], 50)
animate.add("position", "fio", [0, 15], [5, 40], [0, .06, .2, .99], 50)
animate.add("opacity", "fio", 0, 1, [.57, .06, .61, .86], 60)
animate.add("position", "crbutton", [0, 15], [5, 250], [0, .06, .2, .99], 50)

dpg.create_viewport(title='1', width=420, height=380, resizable=True, x_pos=400, y_pos=400)

with dpg.file_dialog(label="Выберите файл",
                     directory_selector=False,
                     show=False,
                     callback=filedialog_ok_callback,
                     tag="file_dialog_id",
                     default_path='',
                     default_filename='',
                     width=500,
                     height=400,
                     modal=True,
                     cancel_callback=cancel_callback_filedialog):
    dpg.add_file_extension("Source files (*.xlsx *.xls){.xlsx,.xls}", color=(0, 255, 255, 255))



dpg.setup_dearpygui()
dpg.show_viewport()
# dpg.start_dearpygui()
dpg.set_primary_window("primeW", True)
while dpg.is_dearpygui_running():
    animate.run()
    dpg.render_dearpygui_frame()

# dpg.set_exit_callback(exit_callback)
dpg.destroy_context()